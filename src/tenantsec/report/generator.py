from __future__ import annotations
import json, datetime, html
from typing import Dict, Any, List, Tuple
from tenantsec.review.scanner import load_sheets_for_ai, run_all_checks
from tenantsec.ai.client import generate_exec_summary, generate_technical_report_md
from tenantsec.ui.templates import get_css  # <-- you already added this

# ---- Optional: DOCX export (install python-docx) ----
try:
    from docx import Document
    from docx.shared import Pt  # , Inches  # <-- Inches unused; remove to avoid warnings
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


def _now_str() -> str:
    return datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")


def _split_md_sections(md: str) -> List[Tuple[str, str]]:
    """
    Very light MD section splitter:
      Returns list of (heading, body) by h1/h2 markers.
      Keeps order. Good enough for our generated technical report.
    """
    lines = md.splitlines()
    sections: List[Tuple[str,str]] = []
    cur_h = "Report"
    buf: List[str] = []
    def flush():
        nonlocal sections, buf, cur_h
        if buf:
            sections.append((cur_h, "\n".join(buf).strip()))
            buf = []
    for ln in lines:
        if ln.startswith("# "):
            flush(); cur_h = ln[2:].strip(); continue
        if ln.startswith("## "):
            flush(); cur_h = ln[3:].strip(); continue
        buf.append(ln)
    flush()
    return sections


def _sev_color(sev: str) -> str:
    s = (sev or "").lower()
    return {
        "critical": "#b71c1c",
        "high": "#d32f2f",
        "medium": "#f57c00",
        "low": "#1976d2",
        "info": "#2e7d32",
    }.get(s, "#444444")


def _md_to_html(md_text: str) -> str:
    """Prefer the markdown lib; fall back to a tiny converter."""
    try:
        import markdown  # pip install markdown
        return markdown.markdown(md_text, extensions=["extra", "sane_lists"])
    except Exception:
        # simple, safe fallback
        esc = html.escape(md_text)
        # very light bold/italic – not perfect, but OK as fallback
        esc = esc.replace("**", "<b>").replace("__", "<i>")
        return "<p>" + esc.replace("\n\n", "</p><p>").replace("\n", "<br/>") + "</p>"


def build_html_report(
    tenant_name: str,
    tenant_id: str,
    exec_json: Dict[str, Any],
    tech_md: str,
    *,
    theme: str = "default",   # <-- allow theme selection
) -> str:
    score = exec_json.get("overall_score", 0)
    risks  = exec_json.get("headline_risks", []) or []
    qwins  = exec_json.get("quick_wins", []) or []
    road   = exec_json.get("roadmap", []) or []

    sections = _split_md_sections(tech_md)

    def h(s): return html.escape(str(s))
    def table(items: List[Dict[str, Any]], columns: List[Tuple[str,str]]) -> str:
        if not items: return "<p><em>None</em></p>"
        th = "".join(f"<th>{h(lbl)}</th>" for key,lbl in columns)
        rows = []
        for it in items:
            tds = "".join(f"<td>{h(it.get(key,''))}</td>" for key,_ in columns)
            rows.append(f"<tr>{tds}</tr>")
        return f"<table class='grid'><thead><tr>{th}</tr></thead><tbody>{''.join(rows)}</tbody></table>"

    # use themed CSS from templates (with internal fallback)
    css = get_css(theme)

    # Headline risks table data
    risk_cols = [("id","Finding ID"),("why","Why"),("impact","Impact"),("priority","Priority")]
    qwin_cols = [("action","Action"),("owner","Owner"),("eta_days","ETA (days)")]

    parts = []
    parts.append(f"""
    <div class="cover">
      <h1>Tenant Security Assessment</h1>
      <div class="small">Generated: {_now_str()}</div>
      <p><strong>Tenant:</strong> {h(tenant_name)} <span class="small">({h(tenant_id)})</span></p>
      <div class="kpi"><strong>Score:</strong> {int(score)}/100</div>
      <div class="kpi"><strong>Users:</strong> {exec_json.get("tenant_meta",{}).get("user_count","—")}</div>
    </div>
    """)

    parts.append("<h2>Executive Summary</h2>")
    parts.append("<h3>Headline Risks</h3>")
    parts.append(table(risks, risk_cols))
    parts.append("<h3>Quick Wins</h3>")
    parts.append(table(qwins, qwin_cols))
    if road:
        parts.append("<h3>Roadmap</h3>")
        for r in road:
            parts.append(f"<p><strong>{h(r.get('theme','Theme'))}</strong></p>")
            items = r.get("items") or []
            parts.append(table(items, [("action","Action"),("eta_days","ETA (days)")]))

    parts.append("<hr/>")
    parts.append("<h2>Technical Remediation Report</h2>")

    # render each MD section with the markdown lib (or fallback)
    for title, body in sections:
        parts.append(f"<h3>{h(title)}</h3>")
        parts.append(_md_to_html(body))

    return f"<!doctype html><html><head><meta charset='utf-8'><style>{css}</style></head><body>{''.join(parts)}</body></html>"


# ----------------------
# DOCX export (optional)
# ----------------------
def _docx_add_title(doc, text: str, size=20, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _docx_add_h2(doc, text: str):
    doc.add_heading(text, level=2)

def _docx_add_para(doc, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)

def _docx_add_table(doc, rows: List[Dict[str, Any]], columns: List[Tuple[str,str]]):
    if not rows:
        _docx_add_para(doc, "None")
        return
    tbl = doc.add_table(rows=1, cols=len(columns))
    hdr = tbl.rows[0].cells
    for i, (_, label) in enumerate(columns):
        hdr[i].text = str(label)
    for r in rows:
        row = tbl.add_row().cells
        for i, (key, _) in enumerate(columns):
            row[i].text = "" if r.get(key) is None else str(r.get(key))

def build_docx_report(path: str, tenant_name: str, tenant_id: str, exec_json: Dict[str, Any], tech_md: str):
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed. Install `python-docx` to export DOCX.")
    doc = Document()

    _docx_add_title(doc, "Tenant Security Assessment", size=26)
    _docx_add_para(doc, f"Tenant: {tenant_name} ({tenant_id})")
    _docx_add_para(doc, f"Generated: {_now_str()}")
    _docx_add_para(doc, f"Score: {exec_json.get('overall_score', 0)}/100")

    _docx_add_h2(doc, "Executive Summary")
    risks  = exec_json.get("headline_risks", []) or []
    qwins  = exec_json.get("quick_wins", []) or []
    road   = exec_json.get("roadmap", []) or []

    _docx_add_para(doc, "Headline Risks")
    _docx_add_table(doc, risks, [("id","Finding ID"),("why","Why"),("impact","Impact"),("priority","Priority")])

    _docx_add_para(doc, "Quick Wins")
    _docx_add_table(doc, qwins, [("action","Action"),("owner","Owner"),("eta_days","ETA (days)")])

    if road:
        _docx_add_para(doc, "Roadmap")
        for r in road:
            _docx_add_para(doc, f"Theme: {r.get('theme','Theme')}")
            _docx_add_table(doc, r.get("items") or [], [("action","Action"),("eta_days","ETA (days)")])

    _docx_add_h2(doc, "Technical Remediation Report")
    sections = _split_md_sections(tech_md)
    for title, body in sections:
        doc.add_heading(title, level=3)
        for line in body.splitlines():
            if not line.strip():
                continue
            if line.startswith(("- ", "* ")):
                _docx_add_para(doc, "• " + line[2:].strip())
            else:
                _docx_add_para(doc, line)

    doc.save(path)


# ----------------------
# Public task
# ----------------------
def generate_reports(tenant_id: str) -> Tuple[Dict[str, Any], str]:
    """
    Returns (exec_summary_json, tech_report_markdown).
    """
    sheets = load_sheets_for_ai(tenant_id)
    findings = run_all_checks(tenant_id)
    exec_json = generate_exec_summary(tenant_id, sheets, findings)
    tech_md   = generate_technical_report_md(tenant_id, sheets, findings)
    # include meta for convenience
    exec_json["tenant_meta"] = exec_json.get("tenant_meta") or {"tenant_id": tenant_id}
    return exec_json, tech_md
